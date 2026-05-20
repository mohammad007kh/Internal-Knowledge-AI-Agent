"""File-upload connector — ingests binary objects stored in MinIO (T-047).

Supported file types
--------------------
* ``pdf``      — structure-aware extraction via ``unstructured`` (pdfminer.six
                 backend, ``strategy='fast'``).  Falls back to PyPDF2 page-text
                 dump on parser failure.
* ``docx``     — structure-aware extraction via ``unstructured`` — preserves
                 tables (as HTML), headings, headers, footers and list
                 bullets.  Falls back to python-docx paragraph-only dump on
                 parser failure.
* ``xlsx``     — plain-text representation (CSV-like) via openpyxl
* ``csv``      — decoded as UTF-8 text
* ``txt``      — decoded as UTF-8 text
* ``md`` /
  ``markdown`` — decoded as UTF-8 Markdown text

Config keys expected in the connector ``config`` dict
------------------------------------------------------

Multi-file shape (preferred)::

    {
        "minio_bucket": "knowledge-agent",
        "files": [
            {
                "object_key":    "uploads/2026/04/...-report.pdf",
                "original_name": "report.pdf",
                "file_type":     "pdf",
                "size_bytes":    12345,
            },
            ...
        ],
        "source_id": "3fa85f64-...",
    }

Legacy single-file shape (still accepted for backward compatibility)::

    {
        "minio_bucket": "knowledge-agent",
        "object_key":   "uploads/abc123/report.pdf",
        "file_type":    "pdf",
        "source_id":    "3fa85f64-...",
    }
"""
from __future__ import annotations

import logging
import uuid
from collections.abc import AsyncIterator
from typing import TYPE_CHECKING, Any

from src.connectors.base import BaseConnector, Document
from src.connectors.registry import register
from src.core.app_config import get_app_config
from src.models.enums import SourceType
from src.schemas.raw_document import RawDocument
from src.services.storage_service import StorageService

if TYPE_CHECKING:
    pass

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Supported file-type constants
# ---------------------------------------------------------------------------

# Keys are the canonical extensions actually consumed by the parsers.
_SUPPORTED_TYPES: frozenset[str] = frozenset({"pdf", "docx", "xlsx", "csv", "txt", "md"})

# Aliases the wizard / API may send — collapsed to the canonical key.
_TYPE_ALIASES: dict[str, str] = {
    "markdown": "md",
}

_TEXT_TYPES: frozenset[str] = frozenset({"csv", "txt", "md"})

_MIME_MAP: dict[str, str] = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "csv": "text/csv",
    "txt": "text/plain",
    "md": "text/plain",
}


def _canonicalise_file_type(file_type: str) -> str:
    """Lower-case, strip a leading dot, and resolve aliases (e.g. markdown→md)."""
    cleaned = file_type.lower().lstrip(".")
    return _TYPE_ALIASES.get(cleaned, cleaned)


# ---------------------------------------------------------------------------
# Connector
# ---------------------------------------------------------------------------


@register(SourceType.FILE_UPLOAD)
class FileUploadConnector(BaseConnector):
    """Ingest one or more files previously uploaded to MinIO.

    The connector accepts either a multi-file ``files`` list (preferred) or
    the legacy single ``object_key`` shape for back-compat.  Each entry is
    downloaded lazily during :meth:`extract_documents` and yielded as a
    separate :class:`~src.connectors.base.Document`.
    """

    def __init__(self, config: dict[str, Any]) -> None:
        super().__init__(config)

        self._bucket: str = config["minio_bucket"]
        self._source_id: str = str(config.get("source_id", ""))
        self._storage: StorageService = StorageService()

        self._files: list[dict[str, Any]] = self._coerce_files_config(config)
        if not self._files:
            raise ValueError(
                "FileUploadConnector requires either 'files' (list) or "
                "'object_key' + 'file_type' in its config."
            )

        for entry in self._files:
            ftype = entry["file_type"]
            if ftype not in _SUPPORTED_TYPES:
                raise ValueError(
                    f"Unsupported file_type '{ftype}'. "
                    f"Supported: {sorted(_SUPPORTED_TYPES)}"
                )

        # Maintain legacy single-file attributes so existing tests keep working.
        self._object_key: str = self._files[0]["object_key"]
        self._file_type: str = self._files[0]["file_type"]
        self._raw_data: bytes | None = None  # bytes for the *first* file (legacy)

    # ------------------------------------------------------------------
    # Config normalisation
    # ------------------------------------------------------------------

    @staticmethod
    def _coerce_files_config(config: dict[str, Any]) -> list[dict[str, Any]]:
        """Return the list of files regardless of which config shape was given."""
        raw_files = config.get("files")
        if isinstance(raw_files, list) and raw_files:
            normalised: list[dict[str, Any]] = []
            for entry in raw_files:
                if not isinstance(entry, dict):
                    raise ValueError(
                        "Each entry in 'files' must be a dict with 'object_key' "
                        "and 'file_type'."
                    )
                object_key = entry.get("object_key")
                file_type = entry.get("file_type")
                if not object_key or not file_type:
                    raise ValueError(
                        "File entry missing required 'object_key' or 'file_type'."
                    )
                normalised.append(
                    {
                        "object_key": str(object_key),
                        "original_name": str(
                            entry.get("original_name")
                            or str(object_key).rsplit("/", 1)[-1]
                        ),
                        "file_type": _canonicalise_file_type(str(file_type)),
                        "size_bytes": entry.get("size_bytes"),
                    }
                )
            return normalised

        # Legacy single-file shape.
        object_key = config.get("object_key")
        file_type = config.get("file_type")
        if object_key and file_type:
            return [
                {
                    "object_key": str(object_key),
                    "original_name": str(object_key).rsplit("/", 1)[-1],
                    "file_type": _canonicalise_file_type(str(file_type)),
                    "size_bytes": None,
                }
            ]
        return []

    # ------------------------------------------------------------------
    # Lifecycle
    # ------------------------------------------------------------------

    async def connect(self) -> None:
        """Download the first object and validate its size + MIME.

        Multi-file downloads are deferred to :meth:`extract_documents` /
        :meth:`fetch_documents` so each file is only held in memory while
        being parsed.  The single-file legacy contract pre-downloads the
        bytes here so existing call sites (and tests) keep working.
        """
        first = self._files[0]
        self._raw_data = await self._download_and_validate(
            object_key=first["object_key"], file_type=first["file_type"]
        )

    async def disconnect(self) -> None:
        """Release any in-memory file bytes."""
        self._raw_data = None

    async def _download_and_validate(self, object_key: str, file_type: str) -> bytes:
        """Download an object from MinIO, enforce size limits + MIME type."""
        app_cfg = get_app_config()
        max_bytes: int = app_cfg.file_upload.max_size_bytes

        raw_data = await self._storage.download_bytes(
            bucket=self._bucket,
            object_key=object_key,
        )

        if len(raw_data) > max_bytes:
            raise ValueError(
                f"File size {len(raw_data):,} bytes exceeds the configured "
                f"maximum of {max_bytes:,} bytes "
                f"(object_key='{object_key}')."
            )

        self._validate_mime_type(raw_data, file_type)

        logger.info(
            "FileUploadConnector: downloaded %d bytes from %s/%s",
            len(raw_data),
            self._bucket,
            object_key,
        )
        return raw_data

    # ------------------------------------------------------------------
    # Document extraction
    # ------------------------------------------------------------------

    async def extract_documents(self) -> AsyncIterator[Document]:  # type: ignore[override]
        """Yield one :class:`Document` per uploaded file.

        For the legacy single-file path the bytes are reused from
        :meth:`connect` (preserves existing test behaviour).  Additional
        files are downloaded lazily here.
        """
        # First file uses pre-downloaded bytes when present.
        first = self._files[0]
        first_bytes = self._raw_data
        if first_bytes is None:
            first_bytes = await self._download_and_validate(
                object_key=first["object_key"], file_type=first["file_type"]
            )
        yield self._build_document(first, first_bytes)

        for entry in self._files[1:]:
            data = await self._download_and_validate(
                object_key=entry["object_key"], file_type=entry["file_type"]
            )
            yield self._build_document(entry, data)

    async def fetch_documents(self) -> list[RawDocument]:
        """Return all uploaded files as :class:`RawDocument` records.

        Used by the Celery sync pipeline (``tasks.sync_source``).
        """
        results: list[RawDocument] = []
        for entry in self._files:
            data = await self._download_and_validate(
                object_key=entry["object_key"], file_type=entry["file_type"]
            )
            text = self._parse_bytes(data, entry["file_type"])
            results.append(
                RawDocument(
                    title=entry["original_name"],
                    url=f"minio://{self._bucket}/{entry['object_key']}",
                    content=text,
                    metadata={
                        "file_type": entry["file_type"],
                        "bucket": self._bucket,
                        "object_key": entry["object_key"],
                        "original_name": entry["original_name"],
                    },
                )
            )
        return results

    def _build_document(self, entry: dict[str, Any], data: bytes) -> Document:
        """Build a :class:`Document` for the given file entry + bytes."""
        text = self._parse_bytes(data, entry["file_type"])
        return Document(
            source_id=uuid.UUID(self._source_id) if self._source_id else uuid.uuid4(),
            raw_text=text,
            raw_storage_path=entry["object_key"],
            metadata={
                "file_type": entry["file_type"],
                "bucket": self._bucket,
                "original_name": entry["original_name"],
            },
        )

    def _parse_bytes(self, data: bytes, file_type: str) -> str:
        """Dispatch to the right parser based on canonical *file_type*."""
        if file_type == "pdf":
            return self._parse_pdf(data)
        if file_type == "docx":
            return self._parse_docx(data)
        if file_type == "xlsx":
            return self._parse_xlsx(data)
        # csv / txt / md — plain UTF-8 text
        return self._parse_text(data)

    # ------------------------------------------------------------------
    # Connection health check
    # ------------------------------------------------------------------

    async def test_connection(self) -> bool:
        """Return *True* when every configured object exists in MinIO."""
        try:
            for entry in self._files:
                exists = await self._storage.object_exists(
                    bucket=self._bucket,
                    object_key=entry["object_key"],
                )
                if not exists:
                    return False
            return True
        except Exception:  # noqa: BLE001
            logger.exception(
                "FileUploadConnector.test_connection() raised unexpectedly "
                "for bucket=%s",
                self._bucket,
            )
            return False

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _validate_mime_type(data: bytes, file_type: str) -> None:
        """Validate that the file magic bytes match the declared extension.

        Raises
        ------
        ValueError
            If the detected MIME type does not match the expected type for
            *file_type*.
        """
        import magic  # noqa: PLC0415

        expected_mime = _MIME_MAP.get(file_type)
        if expected_mime is None:
            return  # unknown type — already blocked by _SUPPORTED_TYPES check

        detected_mime = magic.from_buffer(data[:2048], mime=True)

        # For text-based types (csv, txt, md) both "text/plain" and "text/csv"
        # are acceptable since libmagic often reports "text/plain" for CSV.
        text_mimes = {"text/plain", "text/csv"}
        if expected_mime in text_mimes:
            if detected_mime not in text_mimes:
                raise ValueError(
                    f"File content does not match declared type '{file_type}'. "
                    f"Detected MIME: '{detected_mime}', expected one of {sorted(text_mimes)}."
                )
        elif detected_mime != expected_mime:
            raise ValueError(
                f"File content does not match declared type '{file_type}'. "
                f"Detected MIME: '{detected_mime}', expected '{expected_mime}'."
            )

    # ------------------------------------------------------------------
    # Private parsers
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_pdf(data: bytes) -> str:
        """Extract structured text from PDF bytes via ``unstructured``.

        Uses ``strategy='fast'`` which routes through pdfminer.six — preserves
        reading order (incl. multi-column layouts) and surfaces table cells
        as individual elements.  No OCR is performed; scanned / image-only
        PDFs are still effectively unparseable here.

        Falls back to the legacy PyPDF2 page-text dump on any failure so a
        single bad file cannot break the ingestion pipeline.
        """
        import io  # noqa: PLC0415

        try:
            from unstructured.partition.pdf import partition_pdf  # noqa: PLC0415
            from unstructured.staging.base import elements_to_md  # noqa: PLC0415

            elements = partition_pdf(file=io.BytesIO(data), strategy="fast")
            rendered = elements_to_md(elements)
            if rendered.strip():
                return rendered
        except Exception:  # noqa: BLE001
            logger.warning(
                "FileUploadConnector: unstructured PDF parser failed; "
                "falling back to PyPDF2 page-text dump.",
                exc_info=True,
            )

        import PyPDF2  # type: ignore[import-untyped]  # noqa: PLC0415

        reader = PyPDF2.PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            page_text: str = page.extract_text() or ""
            if page_text.strip():
                parts.append(page_text)
        return "\n".join(parts)

    @staticmethod
    def _parse_docx(data: bytes) -> str:
        """Extract structured text from DOCX bytes via ``unstructured``.

        Captures content python-docx alone could not see: tables (rendered
        as HTML), headers & footers, list bullets, and headings (rendered
        as ``# Heading``).  Output is Markdown / Markdown-with-HTML-tables,
        which is well-tolerated by the downstream chunker and LLM.

        Falls back to the legacy paragraph-only dump on any failure so a
        single bad file cannot break the ingestion pipeline.
        """
        import io  # noqa: PLC0415

        try:
            from unstructured.partition.docx import partition_docx  # noqa: PLC0415
            from unstructured.staging.base import elements_to_md  # noqa: PLC0415

            elements = partition_docx(file=io.BytesIO(data))
            rendered = elements_to_md(elements)
            if rendered.strip():
                return rendered
        except Exception:  # noqa: BLE001
            logger.warning(
                "FileUploadConnector: unstructured DOCX parser failed; "
                "falling back to python-docx paragraph dump.",
                exc_info=True,
            )

        import docx  # type: ignore[import-untyped]  # noqa: PLC0415

        document = docx.Document(io.BytesIO(data))
        paragraphs: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)

    @staticmethod
    def _parse_xlsx(data: bytes) -> str:
        """Extract text from XLSX bytes using openpyxl."""
        import io  # noqa: PLC0415

        import openpyxl  # type: ignore[import-untyped]  # noqa: PLC0415

        workbook = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        rows: list[str] = []
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = ",".join("" if v is None else str(v) for v in row)
                if row_text.strip(","):
                    rows.append(row_text)
        return "\n".join(rows)

    @staticmethod
    def _parse_text(data: bytes) -> str:
        """Decode plain-text bytes (csv / txt / md) as UTF-8."""
        return data.decode("utf-8", errors="replace")
