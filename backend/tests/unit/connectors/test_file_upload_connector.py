"""Unit tests for FileUploadConnector (T-057)."""
from __future__ import annotations

import uuid
from typing import Any
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from src.connectors.base import Document
from src.connectors.file_upload_connector import FileUploadConnector

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_SOURCE_ID = str(uuid.uuid4())
_BUCKET = "uploads"
_OBJECT_KEY = "docs/test.txt"

_TXT_CONTENT = b"This is a plain text document for testing purposes."
_CSV_CONTENT = b"name,value\nalice,1\nbob,2\n"
_MD_CONTENT = b"# Heading\n\nSome **markdown** body."


def _make_connector(
    file_type: str = "txt",
    extra: dict[str, Any] | None = None,
    txt_data: bytes = _TXT_CONTENT,
) -> tuple[FileUploadConnector, MagicMock]:
    """Return (connector, mock_storage_instance) with StorageService patched."""
    cfg: dict[str, Any] = {
        "minio_bucket": _BUCKET,
        "object_key": _OBJECT_KEY,
        "file_type": file_type,
        "source_id": _SOURCE_ID,
    }
    if extra:
        cfg.update(extra)

    with patch("src.connectors.file_upload_connector.StorageService") as mock_cls:
        conn = FileUploadConnector(config=cfg)
        mock_storage = mock_cls.return_value
        conn._storage = mock_storage  # noqa: SLF001

    return conn, mock_storage


def _fake_app_config(max_size_bytes: int = 50 * 1024 * 1024) -> MagicMock:
    cfg = MagicMock()
    cfg.file_upload.max_size_bytes = max_size_bytes
    return cfg


# ---------------------------------------------------------------------------
# Construction — unsupported file type
# ---------------------------------------------------------------------------


def test_constructor_raises_for_unsupported_file_type() -> None:
    with patch("src.connectors.file_upload_connector.StorageService"):
        with pytest.raises(ValueError, match="[Uu]nsupported"):
            FileUploadConnector(config={
                "minio_bucket": _BUCKET,
                "object_key": _OBJECT_KEY,
                "file_type": "exe",
                "source_id": _SOURCE_ID,
            })


def test_constructor_accepts_txt() -> None:
    conn, _ = _make_connector(file_type="txt")
    assert conn is not None


def test_constructor_accepts_csv() -> None:
    conn, _ = _make_connector(file_type="csv")
    assert conn is not None


def test_constructor_accepts_md() -> None:
    conn, _ = _make_connector(file_type="md")
    assert conn is not None


def test_constructor_accepts_pdf() -> None:
    # PDF constructor should not raise even though full parsing is complex;
    # the unsupported-type guard must pass for 'pdf'
    conn, _ = _make_connector(file_type="pdf")
    assert conn is not None


# ---------------------------------------------------------------------------
# connect — happy path
# ---------------------------------------------------------------------------


async def test_connect_downloads_file() -> None:
    conn, mock_storage = _make_connector()
    mock_storage.download_bytes = AsyncMock(return_value=_TXT_CONTENT)

    with patch("src.connectors.file_upload_connector.get_app_config", return_value=_fake_app_config()):
        await conn.connect()

    mock_storage.download_bytes.assert_awaited_once()
    assert conn._raw_data == _TXT_CONTENT  # noqa: SLF001


async def test_connect_raises_value_error_when_file_too_large() -> None:
    conn, mock_storage = _make_connector()
    # Return data larger than max
    large_data = b"x" * 100
    mock_storage.download_bytes = AsyncMock(return_value=large_data)

    cfg = _fake_app_config(max_size_bytes=10)
    with patch("src.connectors.file_upload_connector.get_app_config", return_value=cfg):
        with pytest.raises(ValueError):
            await conn.connect()


# ---------------------------------------------------------------------------
# disconnect
# ---------------------------------------------------------------------------


async def test_disconnect_clears_raw_data() -> None:
    conn, mock_storage = _make_connector()
    mock_storage.download_bytes = AsyncMock(return_value=_TXT_CONTENT)

    with patch("src.connectors.file_upload_connector.get_app_config", return_value=_fake_app_config()):
        await conn.connect()

    await conn.disconnect()
    assert conn._raw_data is None  # noqa: SLF001


# ---------------------------------------------------------------------------
# extract_documents — guard: must call connect first
# ---------------------------------------------------------------------------


async def test_extract_documents_raises_runtime_error_without_connect() -> None:
    conn, _ = _make_connector()
    with pytest.raises(RuntimeError):
        async for _ in conn.extract_documents():
            pass


# ---------------------------------------------------------------------------
# extract_documents — txt happy path
# ---------------------------------------------------------------------------


async def test_extract_documents_txt_yields_one_document() -> None:
    conn, mock_storage = _make_connector(file_type="txt")
    mock_storage.download_bytes = AsyncMock(return_value=_TXT_CONTENT)

    with patch("src.connectors.file_upload_connector.get_app_config", return_value=_fake_app_config()):
        await conn.connect()

    docs: list[Document] = []
    async for doc in conn.extract_documents():
        docs.append(doc)

    assert len(docs) == 1
    assert isinstance(docs[0], Document)


async def test_extract_documents_txt_raw_text() -> None:
    conn, mock_storage = _make_connector(file_type="txt")
    mock_storage.download_bytes = AsyncMock(return_value=_TXT_CONTENT)

    with patch("src.connectors.file_upload_connector.get_app_config", return_value=_fake_app_config()):
        await conn.connect()

    docs: list[Document] = []
    async for doc in conn.extract_documents():
        docs.append(doc)

    assert "plain text document" in docs[0].raw_text


async def test_extract_documents_txt_metadata_file_type() -> None:
    conn, mock_storage = _make_connector(file_type="txt")
    mock_storage.download_bytes = AsyncMock(return_value=_TXT_CONTENT)

    with patch("src.connectors.file_upload_connector.get_app_config", return_value=_fake_app_config()):
        await conn.connect()

    docs: list[Document] = []
    async for doc in conn.extract_documents():
        docs.append(doc)

    assert docs[0].metadata.get("file_type") == "txt"


async def test_extract_documents_txt_source_id_is_uuid() -> None:
    conn, mock_storage = _make_connector(file_type="txt")
    mock_storage.download_bytes = AsyncMock(return_value=_TXT_CONTENT)

    with patch("src.connectors.file_upload_connector.get_app_config", return_value=_fake_app_config()):
        await conn.connect()

    docs: list[Document] = []
    async for doc in conn.extract_documents():
        docs.append(doc)

    assert isinstance(docs[0].source_id, uuid.UUID)
    assert docs[0].source_id == uuid.UUID(_SOURCE_ID)


# ---------------------------------------------------------------------------
# extract_documents — csv happy path
# ---------------------------------------------------------------------------


async def test_extract_documents_csv_yields_one_document() -> None:
    conn, mock_storage = _make_connector(file_type="csv", extra={"object_key": "docs/test.csv"})
    mock_storage.download_bytes = AsyncMock(return_value=_CSV_CONTENT)

    with patch("src.connectors.file_upload_connector.get_app_config", return_value=_fake_app_config()):
        await conn.connect()

    docs: list[Document] = []
    async for doc in conn.extract_documents():
        docs.append(doc)

    assert len(docs) == 1
    assert isinstance(docs[0], Document)


# ---------------------------------------------------------------------------
# extract_documents — markdown happy path
# ---------------------------------------------------------------------------


async def test_extract_documents_md_yields_one_document() -> None:
    conn, mock_storage = _make_connector(file_type="md", extra={"object_key": "docs/readme.md"})
    mock_storage.download_bytes = AsyncMock(return_value=_MD_CONTENT)

    with patch("src.connectors.file_upload_connector.get_app_config", return_value=_fake_app_config()):
        await conn.connect()

    docs: list[Document] = []
    async for doc in conn.extract_documents():
        docs.append(doc)

    assert len(docs) == 1


# ---------------------------------------------------------------------------
# test_connection
# ---------------------------------------------------------------------------


async def test_test_connection_returns_true_when_object_exists() -> None:
    conn, mock_storage = _make_connector()
    mock_storage.object_exists = AsyncMock(return_value=True)

    result = await conn.test_connection()

    assert result is True
    mock_storage.object_exists.assert_called_once()


async def test_test_connection_returns_false_when_object_missing() -> None:
    conn, mock_storage = _make_connector()
    mock_storage.object_exists = AsyncMock(return_value=False)

    result = await conn.test_connection()

    assert result is False


async def test_test_connection_returns_false_on_exception() -> None:
    conn, mock_storage = _make_connector()
    mock_storage.object_exists = AsyncMock(side_effect=Exception("MinIO unreachable"))

    result = await conn.test_connection()

    assert result is False


async def test_test_connection_never_raises() -> None:
    conn, mock_storage = _make_connector()
    mock_storage.object_exists = AsyncMock(side_effect=RuntimeError("unexpected"))

    # Must not propagate
    result = await conn.test_connection()
    assert isinstance(result, bool)
    assert result is False


# ---------------------------------------------------------------------------
# Async context manager
# ---------------------------------------------------------------------------


async def test_async_context_manager_calls_connect_disconnect() -> None:
    conn, _ = _make_connector()

    with patch.object(conn, "connect", new_callable=AsyncMock) as mock_connect, \
         patch.object(conn, "disconnect", new_callable=AsyncMock) as mock_disconnect:
        async with conn:
            mock_connect.assert_called_once()
        mock_disconnect.assert_called_once()


# ---------------------------------------------------------------------------
# Structure-aware DOCX / PDF parsing  (unstructured upgrade)
# ---------------------------------------------------------------------------


def _build_docx_with_table_and_heading() -> bytes:
    """Generate a tiny in-memory DOCX containing a heading + a 3x2 table."""
    import io  # noqa: PLC0415

    import docx  # type: ignore[import-untyped]  # noqa: PLC0415

    document = docx.Document()
    document.add_heading("Project Skills", level=1)
    document.add_paragraph("Resume profile body text.")

    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Years"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "5"
    table.cell(2, 0).text = "Rust"
    table.cell(2, 1).text = "2"

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _build_plain_docx() -> bytes:
    """Generate a paragraph-only DOCX (no tables, no headings)."""
    import io  # noqa: PLC0415

    import docx  # type: ignore[import-untyped]  # noqa: PLC0415

    document = docx.Document()
    document.add_paragraph("Just a plain paragraph for sanity checks.")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _build_pdf_with_table_and_heading() -> bytes:
    """Generate a tiny PDF containing a heading + a 3x2 table via reportlab."""
    import io  # noqa: PLC0415

    from reportlab.lib import colors  # type: ignore[import-untyped]  # noqa: PLC0415
    from reportlab.lib.pagesizes import letter  # type: ignore[import-untyped]  # noqa: PLC0415
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]  # noqa: PLC0415
    from reportlab.platypus import (  # type: ignore[import-untyped]  # noqa: PLC0415
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()

    story = [
        Paragraph("Project Skills Header", styles["Heading1"]),
        Paragraph("Resume profile body text.", styles["BodyText"]),
        Spacer(1, 12),
        Table(
            [["Skill", "Years"], ["Python", "5"], ["Rust", "2"]],
            style=TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black)]),
        ),
    ]
    pdf.build(story)
    return buf.getvalue()


def _build_plain_pdf() -> bytes:
    """Generate a paragraph-only PDF via reportlab."""
    import io  # noqa: PLC0415

    from reportlab.lib.pagesizes import letter  # type: ignore[import-untyped]  # noqa: PLC0415
    from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]  # noqa: PLC0415
    from reportlab.platypus import Paragraph, SimpleDocTemplate  # type: ignore[import-untyped]  # noqa: PLC0415

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    pdf.build([Paragraph("Just a plain paragraph for sanity checks.", styles["BodyText"])])
    return buf.getvalue()


def test_parse_docx_extracts_table_cells() -> None:
    """Real-CV regression: skills/projects in tables must reach the bot."""
    parsed = FileUploadConnector._parse_docx(_build_docx_with_table_and_heading())
    for cell in ("Skill", "Years", "Python", "5", "Rust", "2"):
        assert cell in parsed, f"Table cell '{cell}' missing from parsed DOCX"


def test_parse_docx_extracts_heading_text() -> None:
    parsed = FileUploadConnector._parse_docx(_build_docx_with_table_and_heading())
    assert "Project Skills" in parsed


def test_parse_docx_plain_doc_still_works() -> None:
    """Heading-less / table-less DOCX: legacy happy path must keep working."""
    parsed = FileUploadConnector._parse_docx(_build_plain_docx())
    assert "Just a plain paragraph" in parsed


def test_parse_docx_falls_back_when_unstructured_fails() -> None:
    """If the new parser raises, the legacy python-docx path is used."""
    docx_bytes = _build_plain_docx()

    with patch(
        "unstructured.partition.docx.partition_docx",
        side_effect=RuntimeError("simulated parser failure"),
    ):
        parsed = FileUploadConnector._parse_docx(docx_bytes)

    assert "Just a plain paragraph" in parsed


def test_parse_pdf_extracts_table_cells() -> None:
    parsed = FileUploadConnector._parse_pdf(_build_pdf_with_table_and_heading())
    # Each cell appears as its own element under strategy='fast' — that's fine,
    # the words just need to be retrievable downstream.
    for cell in ("Skill", "Years", "Python", "Rust"):
        assert cell in parsed, f"Table cell '{cell}' missing from parsed PDF"


def test_parse_pdf_extracts_heading_text() -> None:
    parsed = FileUploadConnector._parse_pdf(_build_pdf_with_table_and_heading())
    assert "Project Skills Header" in parsed


def test_parse_pdf_plain_pdf_still_works() -> None:
    parsed = FileUploadConnector._parse_pdf(_build_plain_pdf())
    assert "Just a plain paragraph" in parsed


def test_parse_pdf_falls_back_when_unstructured_fails() -> None:
    """If the new parser raises, the legacy PyPDF2 path is used.

    ``unstructured.partition.pdf`` pulls in heavy native image deps
    (libxcb etc.) that are intentionally absent from the runtime image, so
    the submodule cannot be imported here to patch in place.  We instead
    inject a stand-in module whose ``partition_pdf`` raises, which exercises
    the connector's ``except`` branch exactly as a real parser failure
    would — then assert the PyPDF2 fallback recovered the text.
    """
    import sys  # noqa: PLC0415
    import types  # noqa: PLC0415

    pdf_bytes = _build_plain_pdf()

    fake_module = types.ModuleType("unstructured.partition.pdf")

    def _boom(*_args: Any, **_kwargs: Any) -> None:
        raise RuntimeError("simulated parser failure")

    fake_module.partition_pdf = _boom  # type: ignore[attr-defined]

    with patch.dict(sys.modules, {"unstructured.partition.pdf": fake_module}):
        parsed = FileUploadConnector._parse_pdf(pdf_bytes)

    assert "Just a plain paragraph" in parsed
